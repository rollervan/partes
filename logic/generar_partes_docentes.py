import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- FUNCIONES AUXILIARES ---

def safe_get(row, idx):
    """Obtiene valor seguro de la columna, manejando NaNs."""
    try:
        val = row.iloc[idx]
        if pd.isna(val) or val == "" or str(val).lower() == "nan":
            return "No indicado / No aplica"
        return str(val).strip()
    except IndexError:
        return "N/A"

def add_qa_bloque(doc, pregunta, respuesta):
    """Añade bloque Pregunta (Azul/Negrita) - Respuesta (Normal)."""
    # 1. Pregunta
    p = doc.add_paragraph()
    run = p.add_run(pregunta)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102) # Azul corporativo
    run.font.size = Pt(11)
    
    # 2. Respuesta
    if respuesta and respuesta != "No indicado / No aplica":
        p_res = doc.add_paragraph(respuesta)
        p_res.paragraph_format.left_indent = Inches(0.2) # Sangría visual
        p_res.paragraph_format.space_after = Pt(8)
    else:
        # Texto discreto si no hay respuesta
        p_res = doc.add_paragraph()
        p_res.paragraph_format.left_indent = Inches(0.2)
        run_res = p_res.add_run("Sin comentarios / No aplica.")
        run_res.font.italic = True
        run_res.font.size = Pt(10)
        p_res.paragraph_format.space_after = Pt(8)

# --- FUNCIÓN PRINCIPAL ---

def generar_partes_docentes(df, nombre_salida="Informe_Partes_Docentes.docx"):
    """
    Genera UN ÚNICO documento Word con todas las asignaturas, 
    ordenadas por Curso > Cuatrimestre > Asignatura.
    """
    
    # 1. ORDENAR DATOS (Crucial para el requerimiento)
    # Asumimos columna 9: Curso, 10: Cuatrimestre, 6: Asignatura
    print("🔄 Ordenando datos por Curso y Asignatura...")
    try:
        # Intentamos convertir curso a numérico para ordenar bien (1, 2, 10...), si falla se queda string
        df_sorted = df.copy()
        df_sorted['Curso_Num'] = pd.to_numeric(df_sorted.iloc[:, 9], errors='coerce').fillna(999)
        df_sorted = df_sorted.sort_values(by=['Curso_Num', df.columns[10], df.columns[6]])
    except Exception as e:
        print(f"⚠️ No se pudo ordenar numéricamente, usando orden alfabético: {e}")
        df_sorted = df.sort_values(by=[df.columns[9], df.columns[6]])

    # 2. CREAR DOCUMENTO
    doc = Document()
    
    # Portada General
    titulo_principal = doc.add_heading('INFORME DE CALIDAD DOCENTE', 0)
    titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Total de asignaturas procesadas: {len(df)}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    print(f"🚀 Generando informe único con {len(df_sorted)} asignaturas...")

    # 3. BUCLE DE GENERACIÓN
    for index, row in df_sorted.iterrows():
        
        # --- ENCABEZADO DE ASIGNATURA ---
        asignatura = safe_get(row, 6)
        profesor = safe_get(row, 4)
        curso = safe_get(row, 9)
        
        # Título de la sección (Nombre Asignatura)
        h1 = doc.add_heading(asignatura, level=1)
        h1.style.font.color.rgb = RGBColor(0, 0, 0) # Negro para imprimir bien
        
        # Subtítulo (Metadatos)
        p_meta = doc.add_paragraph()
        p_meta.add_run(f"Curso: {curso} | Cuatrimestre: {safe_get(row, 10)}").bold = True
        p_meta.add_run(f"\nProfesor/a: {profesor}")
        p_meta.add_run(f"\nTitulación: {safe_get(row, 5)}")

        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER # Separador visual

        # --- SECCIÓN 1: DATOS CUANTITATIVOS ---
        doc.add_heading('1. Datos Cuantitativos', level=2)
        
        matriculados = safe_get(row, 8)
        aprobados = safe_get(row, 7)
        
        # Tabla simple para KPIs
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True
        cells = table.rows[0].cells
        
        # Rellenar celdas
        cells[0].text = f"Matriculados: {matriculados}"
        cells[1].text = f"Aprobados: {aprobados}"
        
        # Calcular %
        try:
            pct = (float(aprobados) / float(matriculados)) * 100
            cells[2].text = f"Tasa Éxito: {pct:.1f}%"
        except:
            cells[2].text = "Tasa Éxito: N/A"
            
        # Poner negrita en la tabla
        for c in cells:
            if c.paragraphs: c.paragraphs[0].runs[0].bold = True

        doc.add_paragraph() # Espacio

        # --- SECCIÓN 2: RESULTADOS ---
        doc.add_heading('2. Análisis de Resultados', level=2)
        add_qa_bloque(doc, "Valoración (1-5):", safe_get(row, 12))
        add_qa_bloque(doc, "Justificación / Acciones:", safe_get(row, 13))
        
        if "sí" in safe_get(row, 14).lower():
            add_qa_bloque(doc, "Deficiencias previas:", safe_get(row, 15))

        # --- SECCIÓN 3: DOCENCIA ---
        doc.add_heading('3. Docencia e Incidencias', level=2)
        
        temario = safe_get(row, 18)
        add_qa_bloque(doc, "¿Temario completo?:", temario)
        if "no" in temario.lower():
            add_qa_bloque(doc, "Causa:", safe_get(row, 19))
            
        add_qa_bloque(doc, "Incidencias / Problemas:", 
                      f"{safe_get(row, 20)}\n{safe_get(row, 21)}")
        
        detalles = safe_get(row, 22)
        if detalles != "No indicado / No aplica":
            add_qa_bloque(doc, "Detalles adicionales:", detalles)

        # --- SECCIÓN 4: COORDINACIÓN Y GRUPO ---
        doc.add_heading('4. Grupo y Coordinación', level=2)
        add_qa_bloque(doc, "Características Grupo:", safe_get(row, 23))
        add_qa_bloque(doc, "Satisfacción Grupo:", safe_get(row, 24))
        add_qa_bloque(doc, "Coordinación:", safe_get(row, 30))

        # --- SECCIÓN 5: CIERRE ---
        doc.add_heading('5. Cierre', level=2)
        add_qa_bloque(doc, "Otras incidencias:", safe_get(row, 33))
        add_qa_bloque(doc, "Sugerencias:", safe_get(row, 34))

        # --- FIN DE ASIGNATURA ---
        # Salto de página para la siguiente asignatura (excepto si es la última)
        if index != df_sorted.index[-1]:
            doc.add_page_break()

    # 4. GUARDAR
    try:
        doc.save(nombre_salida)
        print(f"✅ Documento guardado exitosamente: {nombre_salida}")
    except PermissionError:
        print("❌ Error: Cierra el archivo Word si lo tienes abierto e inténtalo de nuevo.")
